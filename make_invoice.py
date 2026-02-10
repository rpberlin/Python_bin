#!/Library/Frameworks/Python.framework/Versions/3.10/bin/python3
from docx import Document
from docx.shared import Inches
from datetime import datetime
import calendar 
from docx2pdf import convert as docx2pdf_convert


print("Enter the month number (1 for January, 2 for February, etc")
month = int(input()) 
print("Enter the year")
year = int(input())

first_day, last_day = calendar.monthrange(year, month)
print(first_day, last_day)

todays_date = datetime.today().strftime('%d-%m-%Y')

Months = ['Januar','Februar','Maerz','April','Mai','Juni','Juli','August','September','Oktober','November','Dezember']


# Create the Word document
doc = Document()

doc.add_heading("Rechnung", level=1)

doc.add_paragraph("Rechnung Nr.: "+str(year)+"-"+str(month))
doc.add_paragraph("Ausstellungsdatum: "+todays_date)

doc.add_heading("Leistender Unternehmer", level=2)
doc.add_paragraph("Ryan Blanchard\nKlosterheider Weg 35 \n13467 Berlin\nDeutschland\nSteuernummer: 00/000/00000")

doc.add_heading("Leistungsempfänger", level=2)
doc.add_paragraph("Oklo Inc\n3190 Coronado Dr\nSanta Clara, CA 95054\nUSA")

doc.add_paragraph(f'Leistungszeitraum: 1–{str(last_day)} {Months[month-1]} {str(year)}')
doc.add_paragraph("Leistungsbeschreibung: Engineering-Consulting (remote)")

# Add invoice table
table = doc.add_table(rows=1, cols=5)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Pos.'
hdr_cells[1].text = 'Beschreibung'
hdr_cells[2].text = 'Menge'
hdr_cells[3].text = 'Einheitspreis (USD)'
hdr_cells[4].text = 'Gesamt (USD)'

row_cells = table.add_row().cells
row_cells[0].text = '1'
row_cells[1].text = 'Consulting-Retainer '+ Months[month-1] +' '+str(year)
row_cells[2].text = '1'
row_cells[3].text = '$17500'
row_cells[4].text = '$17500'

doc.add_paragraph("\nNettobetrag (USD): $17500")
doc.add_paragraph("Umsatzsteuer: 0,00 USD")
doc.add_paragraph("Rechnungsbetrag: $17500")

doc.add_paragraph(
    '„Keine Umsatzsteuer ausgewiesen, da Leistung nicht steuerbar gem. § 3a Abs. 2 UStG – '
    'Ort der sonstigen Leistung: USA (Reverse Charge)“'
)

doc.add_paragraph("Zahlung bereits am "+str(last_day)+"/"+str(month)+"/"+str(year)+" eingegangen.")


word_path = "Invoice_"+Months[month-1]+"_"+str(year)+'-'+str(month)+'.docx'
doc.save(word_path)

docx2pdf_convert(word_path)
